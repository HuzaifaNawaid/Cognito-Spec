import streamlit as st
from ai_controller import ai_controller
from firebase_manager import firebase_manager
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- HIDE DEFAULT NAVIGATION CSS INJECTION ---
st.markdown("""
<style>
/* Targets the automatically generated navigation div and hides it */
div[data-testid="stSidebarNav"] {
    display: none;
}
</style>
""", unsafe_allow_html=True)
# --------------------------------------------------

# Page configuration
st.set_page_config(
    page_title="Cognito-Spec",
    page_icon="logooo.png",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        'Get Help': None,
        'Report a bug': None,
        'About': None
    }
)

# --- CUSTOM SELECTION BOX NAVIGATION LOGIC ---
def navigate_with_selectbox():
    """Handles the sidebar selectbox navigation"""
    
    page_names = {
        "🏠 Home": "Home.py",
        "⚙️ Project Setup": "pages/1_Project_Setup.py",
        "💬 Interview": "pages/2_Interview.py",
        "🔍 Gap Analysis": "pages/3_Gap_Analysis.py",
        "📄 SRS Generation": "pages/4_Export.py"
    }
    
    # Determine current page name
    current_page_name = "🔍 Gap Analysis"
    
    selected_page_name = st.selectbox(
        "**Go to Page:**", 
        options=list(page_names.keys()), 
        index=list(page_names.keys()).index(current_page_name),
        key="custom_sidebar_select"
    )
    
    # Check if the selected page is different from the current page and switch
    if selected_page_name != current_page_name:
        st.session_state.current_page = selected_page_name
        st.switch_page(page_names[selected_page_name])

# Function to create Word document from report content
def create_word_document(report_content, project_name):
    """Convert report content to Word document"""
    try:
        # Create a new Document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'Risk Assessment Report - {project_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_para = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Add empty line
        
        # Process the report content
        lines = report_content.split('\n')
        
        for line in lines:
            if line.startswith('# '):
                # Main heading
                heading = doc.add_heading(line[2:], 1)
            elif line.startswith('## '):
                # Subheading
                heading = doc.add_heading(line[3:], 2)
            elif line.startswith('### '):
                # Sub-subheading
                heading = doc.add_heading(line[4:], 3)
            elif line.startswith('**') and '**' in line:
                # Bold text
                para = doc.add_paragraph()
                run = para.add_run(line.replace('**', ''))
                run.bold = True
            elif line.startswith('- ') or line.startswith('* '):
                # List item
                para = doc.add_paragraph(style='List Bullet')
                para.add_run(line[2:])
            elif line.strip() == '':
                # Empty line
                doc.add_paragraph()
            else:
                # Regular paragraph
                para = doc.add_paragraph(line)
        
        # Save to bytes buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
        
    except Exception as e:
        st.error(f"Error creating Word document: {e}")
        return None

# Initialize session state
def init_session_state():
    if 'project_initialized' not in st.session_state:
        st.session_state.project_initialized = False
    if 'project_id' not in st.session_state:
        st.session_state.project_id = None
    if 'project_name' not in st.session_state:
        st.session_state.project_name = ""
    if 'domain' not in st.session_state:
        st.session_state.domain = "Healthcare"
    if 'gap_analysis_report' not in st.session_state:
        st.session_state.gap_analysis_report = None

init_session_state()

# Sidebar navigation
with st.sidebar:
    try:
        st.image("logooo.png", width=270)
    except:
        st.title("🚀 Cognito-Spec") 
    st.markdown("---")
    
    # Custom Navigation Selectbox
    navigate_with_selectbox()
    
    st.markdown("---")
    
    # Project Info
    if st.session_state.project_initialized:
        st.subheader("Current Project Status")
        
        # Displaying all requested details
        status_icon = "✅ Initialized"
        req_count = len(st.session_state.requirements)
        
        st.markdown(f"**Name:** **`{st.session_state.project_name}`**")
        st.markdown(f"**Domain:** **`{st.session_state.domain}`**")
        st.markdown(f"**Status:** **`{status_icon}`**")
        st.markdown(f"**Requirements:** **`{req_count}`**")

st.title("🔍 Gap Analysis & Validation")

if not st.session_state.project_initialized:
    st.warning("Please initialize your project first.")
    st.switch_page("pages/1_Project_Setup.py")

# Display project info
st.subheader(f"Project: {st.session_state.project_name}")
st.caption(f"Domain: {st.session_state.domain}")

# Gap Analysis Section
st.header("Run Gap Analysis")

st.write("Click the button below to perform a comprehensive gap analysis on your collected requirements.")

if st.button("Run Gap Analysis", type="primary", use_container_width=True):
    with st.spinner("Performing gap analysis..."):
        gap_analysis = ai_controller.perform_gap_analysis(
            st.session_state.project_id,
            st.session_state.domain
        )
        
        # Store the report in session state
        st.session_state.gap_analysis_report = gap_analysis
        
        st.subheader("Risk Assessment Report")
        st.markdown(gap_analysis)

# Display download button if report exists
if st.session_state.gap_analysis_report:
    st.markdown("---")
    st.header("Download Report")
    
    # Create Word document
    word_buffer = create_word_document(
        st.session_state.gap_analysis_report,
        st.session_state.project_name
    )
    
    if word_buffer:
        st.download_button(
            label="📥 Download as Word Document (.docx)",
            data=word_buffer,
            file_name=f"{st.session_state.project_name}_Risk_Assessment_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )