import streamlit as st
from ai_controller import ai_controller
from firebase_manager import firebase_manager
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from datetime import datetime

# --- HIDE DEFAULT NAVIGATION CSS INJECTION ---
st.markdown("""
<style>
/* Targets the automatically generated navigation div and hides it */
div[data-testid="stSidebarNav"] {
    display: none;
}
</style>
""", unsafe_allow_html=True)
# --------------------------------------------------
# Page configuration
st.set_page_config(
    page_title="Cognito-Spec",
    page_icon="logooo.png",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        'Get Help': None,
        'Report a bug': None,
        'About': None
    }
)
# --- CUSTOM SELECTION BOX NAVIGATION LOGIC ---
def navigate_with_selectbox():
    """Handles the sidebar selectbox navigation"""
    
    page_names = {
        "🏠 Home": "Home.py",
        "⚙️ Project Setup": "pages/1_Project_Setup.py",
        "💬 Interview": "pages/2_Interview.py",
        "🔍 Gap Analysis": "pages/3_Gap_Analysis.py",
        "📄 SRS Generation": "pages/4_Export.py"
    }
    
    # Determine current page name
    current_page_name = "📄 SRS Generation"
    
    selected_page_name = st.selectbox(
        "**Go to Page:**", 
        options=list(page_names.keys()), 
        index=list(page_names.keys()).index(current_page_name),
        key="custom_sidebar_select"
    )
    
    # Check if the selected page is different from the current page and switch
    if selected_page_name != current_page_name:
        st.session_state.current_page = selected_page_name
        st.switch_page(page_names[selected_page_name])

# Initialize session state
def init_session_state():
    if 'project_initialized' not in st.session_state:
        st.session_state.project_initialized = False
    if 'project_id' not in st.session_state:
        st.session_state.project_id = None
    if 'project_name' not in st.session_state:
        st.session_state.project_name = ""
    if 'domain' not in st.session_state:
        st.session_state.domain = "Healthcare"
    if 'srs_generated' not in st.session_state:
        st.session_state.srs_generated = False
    if 'srs_document' not in st.session_state:
        st.session_state.srs_document = None

init_session_state()

# Sidebar navigation
with st.sidebar:
    try:
        st.image("logooo.png", width=270)
    except:
        st.title("🚀 Cognito-Spec") 
    st.markdown("---")
    
    # Custom Navigation Selectbox
    navigate_with_selectbox()
    
    st.markdown("---")
    
        # Project Info
    if st.session_state.project_initialized:
        st.subheader("Current Project Status")
        
        # Displaying all requested details
        status_icon = "✅ Initialized"
        req_count = len(st.session_state.requirements)
        
        st.markdown(f"**Name:** **`{st.session_state.project_name}`**")
        st.markdown(f"**Domain:** **`{st.session_state.domain}`**")
        st.markdown(f"**Status:** **`{status_icon}`**")
        st.markdown(f"**Requirements:** **`{req_count}`**")

st.title("📋 Software Requirements Specification")

if not st.session_state.project_initialized:
    st.warning("Please initialize your project first.")
    st.switch_page("pages/1_Project_Setup.py")

# Display project info
st.subheader(f"Project: {st.session_state.project_name}")
st.caption(f"Domain: {st.session_state.domain}")

# Check if we have requirements
if st.session_state.project_id:
    requirements = firebase_manager.fetch_all_requirements(st.session_state.project_id)
    
    if not requirements:
        st.warning("No requirements collected yet. Please visit the Interview page to collect requirements first.")
        st.info("Navigate to Interview page to start collecting requirements.")
    else:
        st.success(f"✅ {len(requirements)} requirements ready for SRS generation")

# SRS Document Generation
st.header("Generate SRS Document")

st.write("Click the button below to compile all collected requirements into a professional Software Requirements Specification document.")

if st.button("Generate SRS Document", type="primary", use_container_width=True):
    with st.spinner("Compiling SRS document..."):
        srs_content = ai_controller.compile_srs_document(st.session_state.project_id)
        
        # Clean the content by removing * and # characters
        cleaned_lines = []
        for line in srs_content.split('\n'):
            # Remove asterisks but keep the text (Markdown artifacts)
            line = line.replace('*', '').replace('#', '')
            # Clean up multiple spaces
            line = ' '.join(line.split())
            cleaned_lines.append(line)
        
        # --- IMPROVED FILTERING LOGIC ---
        # This explicitly ignores EVERYTHING until it sees "1. INTRODUCTION"
        # This ensures Version, Date, Prepared By, and --- are strictly removed.
        content_started = False
        preview_lines = []
        processed_lines = []
        
        for line in cleaned_lines:
            stripped_line = line.strip()
            
            if not content_started:
                # We strictly look for the start of Section 1.
                # Common variations: "1. INTRODUCTION", "1 INTRODUCTION", "1. Introduction"
                if stripped_line.startswith("1.") or (stripped_line.startswith("1 ") and "introduction" in stripped_line.lower()):
                    content_started = True
                    preview_lines.append(line)
                    processed_lines.append(line)
                else:
                    # Skip all header metadata (Version, Date, Prepared By, ---)
                    continue
            else:
                # Once content has started, we keep everything
                preview_lines.append(line)
                processed_lines.append(line)
        
        # Build the preview content that matches the downloadable document
        # We manually inject the clean header here for the preview only
        preview_content = f"""Software Requirements Specification

for

    {st.session_state.project_name}

Version 1.0 approved

Prepared by: Requirements Engineer

Domain: {st.session_state.domain}

Date: {datetime.now().strftime('%d-%m-%Y')}

---

{chr(10).join(preview_lines)}"""
        
        st.subheader("📄 SRS Document Preview")
        st.markdown(preview_content)
        
        # Create Word document
        doc = Document()
        
        # ==================== FRONT PAGE ====================
        # Add title with slightly larger font
        front_page = doc.add_paragraph()
        front_page_run = front_page.add_run("Software Requirements Specification\n")
        front_page_run.font.size = Pt(28)  # Slightly larger than normal
        front_page_run.font.bold = True
        front_page.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        doc.add_paragraph("\n")
        
        # Add "for" line
        for_line = doc.add_paragraph()
        for_run = for_line.add_run("for\n")
        for_run.font.size = Pt(18)  # Slightly larger
        for_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add project name
        project_line = doc.add_paragraph()
        project_run = project_line.add_run(f"    {st.session_state.project_name}\n")
        project_run.font.size = Pt(22)  # Slightly larger
        project_run.bold = True
        project_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        doc.add_paragraph("\n\n")
        
        # Add version and approval info
        version_line = doc.add_paragraph()
        version_run = version_line.add_run("Version 1.0 approved\n\n")
        version_run.font.size = Pt(14)  # Slightly larger
        version_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add prepared by info
        info_lines = [
            "Prepared by: Requirements Engineer",
            f"Domain: {st.session_state.domain}",
            f"Date: {datetime.now().strftime('%d-%m-%Y')}"
        ]
        
        for line in info_lines:
            info_para = doc.add_paragraph()
            info_run = info_para.add_run(f"{line}\n")
            info_run.font.size = Pt(12)  # Slightly larger than normal text
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page break
        doc.add_page_break()
        
        # ==================== MAIN CONTENT ====================
        # Process the cleaned content for the document
        for line in processed_lines:
            if not line.strip():
                # Empty line - add paragraph break
                doc.add_paragraph()
                continue
            
            # Check if it's a heading (looks like "1. INTRODUCTION" or "1.1 Purpose")
            if line.strip() and len(line.strip()) < 100 and any(line.strip().startswith(f"{i}.") for i in range(1, 10)):
                # Determine heading level
                if '.' in line:
                    parts = line.split('.')
                    # Check if it's a main heading (single number with text)
                    if len(parts) == 2 and parts[0].strip().isdigit() and parts[1].strip():
                        # Main section heading (e.g., "1. INTRODUCTION")
                        doc.add_heading(line.strip(), level=1)
                    elif len(parts) == 2 and parts[0].strip().isdigit() and parts[1].strip().isdigit():
                        # Subsection heading (e.g., "1.1 Purpose")
                        doc.add_heading(line.strip(), level=2)
                    elif len(parts) >= 3 and parts[0].strip().isdigit() and parts[1].strip().isdigit():
                        # Sub-subsection heading (e.g., "2.1.1 User Registration")
                        doc.add_heading(line.strip(), level=3)
                    else:
                        doc.add_paragraph(line.strip())
                else:
                    doc.add_paragraph(line.strip())
            else:
                # Regular content
                doc.add_paragraph(line.strip())
        
        # Save to bytes buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        st.session_state.srs_document = buffer
        st.session_state.srs_generated = True

# Sign-off and Export
if st.session_state.get('srs_generated', False):
    st.header("✅ Sign-off & Export")
    
    st.write("Review the SRS document above and confirm that it accurately represents your system needs.")
    
    sign_off = st.checkbox("I confirm that these requirements accurately represent the system needs")
    
    if sign_off:
        st.success("✅ Requirements approved and ready for export!")
        
        st.download_button(
            label="📥 Download SRS Document (.docx)",
            data=st.session_state.srs_document,
            file_name=f"SRS_{st.session_state.project_name.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    else:
        st.info("Please provide sign-off to enable export.")
else:
    st.info("Generate the SRS document first to enable export.")